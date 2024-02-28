# Version 2.2.0
import datetime
import os
import time
import extract
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
import win32api, win32print

def loadFile(filePath: str) -> list[str]:
    with open(filePath, "r", encoding="utf-8")as f:
        contentList = f.readlines()
    return contentList

def loadUseageWay(filePath: str) -> dict:
    useageWayDict = {}
    with open(filePath, "r", encoding="utf-8")as f:
        for data in f.readlines():
            data = data.split("=")
            useageWayDict[data[0].strip()] = data[1].strip()
    return useageWayDict

def loadPrint(filePath: str, printerName: str, receiveNumber: str, ptName: str):
    filePath = os.getcwd() + filePath
    win32api.ShellExecute(
        0,
        "print",
        filePath,
        "/d:{}".format(win32print.OpenPrinter(printerName)),
        ".",
        0
    )
    print("領藥號：{}，病人：{}，列印完成。".format(receiveNumber, ptName))
    pass

def msWordFormat(pageWd: float, pageHt: float, marginL: float, marginR: float, marginT: float, marginB: float) -> docx.Document():
    msDoc = docx.Document()
    section = msDoc.sections[0]
    section.page_width = Cm(pageWd)

    section.page_height = Cm(pageHt)

    section.left_margin = Cm(marginL)

    section.right_margin = Cm(marginR)

    section.top_margin = Cm(marginT)

    section.bottom_margin = Cm(marginB)

    return msDoc

def drugBagMaker(contentList: list[str], useageWayDict: dict, frequencyDict: dict, beforeOrAfterDict: dict, envSettingDict: dict):
    msDoc = msWordFormat(float(envSettingDict["pageWd"]), float(envSettingDict["pageHt"]),
                         float(envSettingDict["marginL"]), float(envSettingDict["marginR"]),
                         float(envSettingDict["marginT"]), float(envSettingDict["marginB"]))
    pharmacistName = envSettingDict["調劑藥師"]
    printerName = envSettingDict["印表機名稱"]

    receiveNumber: str = extract.extractReceiveNumber(contentList)
    ptName: str = extract.extractPtName(contentList)
    ptBirthDay: str = extract.extractBirthDay(contentList)
    dipensingDay: datetime.strftime = datetime.datetime.now().strftime("%Y/%m/%d")
    ptChartNumber: str = extract.extractChartNumber(contentList)
    department = extract.extractDepartment(contentList)
    doctorName = extract.extractDoctorName(contentList)

    drugNameList, usageList, brandNameAndNoticeList = extract.extractMedisonInfo(contentList)

    paragraph_format = msDoc.styles['Normal'].paragraph_format
    paragraph_format.space_after = 1

    drugCount = len(drugNameList)

    for index in range(drugCount):
        headerTable = msDoc.add_table(rows=4, cols=4)
        msDoc.add_paragraph()
        headerTable.alignment = WD_TABLE_ALIGNMENT.RIGHT
        contentTable = msDoc.add_table(rows=5, cols=3)
        contentTable.cell(0, 1).width = Cm(8)


        headerTable.rows[0].cells[3].text = receiveNumber + " 林口"
        headerTable.rows[0].cells[3].paragraphs[0].runs[0].font.bold = True
        headerTable.rows[1].cells[0].text = ptName
        headerTable.rows[1].cells[2].text = ptBirthDay
        headerTable.rows[2].cells[0].text = ptChartNumber
        headerTable.rows[2].cells[3].text = dipensingDay
        headerTable.rows[3].cells[0].text = department
        headerTable.rows[3].cells[1].text = doctorName
        headerTable.rows[3].cells[3].text = pharmacistName

        headerTable.rows[0].cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        headerTable.rows[1].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        headerTable.rows[2].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        headerTable.rows[2].cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        headerTable.rows[3].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        headerTable.rows[3].cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


        contentTable.rows[0].cells[0].text = chr(12304) + "藥名" + chr(12305)
        contentTable.rows[0].cells[1].text = drugNameList[index]
        contentTable.rows[0].cells[2].text = "{} PC".format(usageList[index][-2].replace("PC", ""))
        contentTable.rows[1].cells[0].text = chr(12304) + "商品名" + chr(12305)
        contentTable.rows[1].cells[1].text = brandNameAndNoticeList[index][0]
        contentTable.rows[2].cells[0].text = chr(12304) + "使用方法" + chr(12305)
        contentTable.rows[2].cells[1].text = "{}".format(useageWayDict.get(usageList[index][2], "None"))
        contentTable.rows[2].cells[2].text = "{} - {}".format(index + 1, drugCount)
        contentTable.rows[3].cells[1].text = "每次{}，{}，{}".format(
                                    usageList[index][0],
                                    frequencyDict.get(usageList[index][1], "None"),
                                    beforeOrAfterDict.get(usageList[index][3], ""))
        contentTable.rows[4].cells[0].text = chr(12304) + "備註" + chr(12305)
        contentTable.rows[4].cells[1].text = brandNameAndNoticeList[index][1]

        for row in contentTable.rows:
            row.cells[-1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        if index != drugCount - 1:
            msDoc.add_page_break()

    msDoc.save("./history/{}_{}.docx".format(receiveNumber, dipensingDay.replace("/", "")))
    print("領藥號：{}，病人：{}，已完成。".format(receiveNumber, ptName))
    loadPrint("/history/{}_{}.docx".format(receiveNumber, dipensingDay.replace("/", "")), printerName, receiveNumber, ptName)
    pass

def envSet(filePath: str) -> dict:
    with open(filePath, "r", encoding="utf-8")as f:
        envSettingDict = {}
        for setting in f.readlines():
            setting = setting.split("=")
            envSettingDict[setting[0].strip()] = setting[1].strip()

    return envSettingDict

if __name__ == "__main__":
    useageWayDict = loadUseageWay("./使用方式.txt")
    frequencyDict = loadUseageWay("./頻次.txt")
    envSettingDict = envSet("./env")
    beforeOrAfterDict = {"PC": "飯後", "AC": "飯前"}
    print("製作人員：謝昀燊Vincent")
    print("Version：2.2.0")
    while True:
        filePath = input("請輸入文字檔路徑：")
        contentList = loadFile(filePath)
        drugBagMaker(contentList, useageWayDict, frequencyDict, beforeOrAfterDict, envSettingDict)
    pass
